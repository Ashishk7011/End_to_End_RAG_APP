from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Iterable

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx"}


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_pdf(path: Path) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    docs: list[Document] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": path.name, "page": page_number},
                )
            )
    return docs


def _read_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def load_uploaded_files(uploaded_files: Iterable) -> list[Document]:
    docs: list[Document] = []
    for uploaded_file in uploaded_files:
        suffix = Path(uploaded_file.name).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            continue

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getbuffer())
            tmp_path = Path(tmp.name)

        try:
            if suffix == ".pdf":
                docs.extend(_read_pdf(tmp_path))
            elif suffix == ".docx":
                docs.append(
                    Document(
                        page_content=_read_docx(tmp_path),
                        metadata={"source": uploaded_file.name},
                    )
                )
            else:
                docs.append(
                    Document(
                        page_content=_read_txt(tmp_path),
                        metadata={"source": uploaded_file.name},
                    )
                )
        finally:
            tmp_path.unlink(missing_ok=True)

    return [doc for doc in docs if doc.page_content.strip()]


def split_documents(
    documents: list[Document],
    *,
    chunk_size: int,
    chunk_overlap: int,
) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = index
    return chunks
